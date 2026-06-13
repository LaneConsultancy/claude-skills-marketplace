#!/usr/bin/env python3
"""
ServiceM8 Form Template Generator

Generates professional DOCX templates from SM8F form files with full customisation.

Usage:
    python3 generate_form.py \
        --sm8f "form.sm8f" \
        --company "Company Name" \
        --logo "logo.png" \
        --primary "#004aad" \
        --style bold_trade \
        --layout compact \
        --density tight \
        --repeat-fields grid \
        --orientation portrait \
        --output template.docx
"""

import argparse
import json
import zipfile
import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    color = color_hex.replace('#', '')
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, border_color="BFBFBF", width="4"):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{width}" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="{width}" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="{width}" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="{width}" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def set_row_height(row, height_pt):
    """Set minimum row height"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="atLeast"/>')
    trPr.append(trHeight)

def add_merge_field(paragraph, field_name, font_size=8, bold=False, font_name='Arial'):
    """Add a Word merge field"""
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    
    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> MERGEFIELD {field_name} \\* MERGEFORMAT </w:instrText>')
    run2._r.append(instrText)
    
    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)
    
    run4 = paragraph.add_run(f"«{field_name}»")
    run4.font.size = Pt(font_size)
    run4.font.name = font_name
    run4.font.bold = bold
    
    run5 = paragraph.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

def darken_color(hex_color, factor=0.7):
    """Darken a hex color"""
    hex_color = hex_color.replace('#', '')
    r = int(int(hex_color[0:2], 16) * factor)
    g = int(int(hex_color[2:4], 16) * factor)
    b = int(int(hex_color[4:6], 16) * factor)
    return f"#{r:02x}{g:02x}{b:02x}"

def name_to_merge_field(name):
    """Convert field name to merge field format"""
    clean = re.sub(r'[^a-zA-Z0-9\s]', '', name.lower())
    clean = re.sub(r'\s+', '_', clean.strip())
    return f"form_{clean}"


# =============================================================================
# FORM PARSING
# =============================================================================

def extract_form(sm8f_path):
    """Extract and parse SM8F form file"""
    with zipfile.ZipFile(sm8f_path, 'r') as z:
        with z.open('form.json') as f:
            data = json.load(f)
    
    form_name = data['form']['name']
    fields = []
    
    for field in sorted(data['fields'], key=lambda x: int(x.get('sort_order', 0))):
        field_data = json.loads(field['field_data_json'])
        fields.append({
            'uuid': field['uuid'],
            'name': field['name'],
            'type': field_data.get('fieldType', 'Text'),
            'choices': field_data.get('choices', []),
            'mandatory': field_data.get('mandatory', False),
            'conditions': field_data.get('conditions', []),
            'merge_field': name_to_merge_field(field['name'])
        })
    
    return form_name, fields

def detect_repeating_groups(fields):
    """Detect numbered repeating field groups (e.g., Appliance 1, Appliance 2)"""
    groups = {}
    pattern = re.compile(r'^(.+?)[\s_]?(\d+)$')
    
    for field in fields:
        match = pattern.match(field['name'])
        if match:
            base_name = match.group(1).strip()
            number = int(match.group(2))
            if base_name not in groups:
                groups[base_name] = {'numbers': set(), 'fields': []}
            groups[base_name]['numbers'].add(number)
            groups[base_name]['fields'].append(field)
    
    # Only return groups with multiple instances
    repeating = {k: v for k, v in groups.items() if len(v['numbers']) > 1}
    return repeating


# =============================================================================
# CONFIGURATION
# =============================================================================

class Config:
    """Template configuration"""
    def __init__(self, args):
        self.company = args.company
        self.logo = args.logo
        self.primary_color = args.primary
        self.secondary_color = darken_color(args.primary)
        self.style = args.style
        self.layout = args.layout
        self.density = args.density
        self.repeat_fields = args.repeat_fields
        self.orientation = args.orientation
        
        # Density settings
        density_map = {
            'tight': {'font': 7, 'label': 7, 'title': 12, 'row': 14, 'margin': 0.3, 'spacing': 2},
            'normal': {'font': 8, 'label': 8, 'title': 14, 'row': 18, 'margin': 0.5, 'spacing': 6},
            'spacious': {'font': 9, 'label': 9, 'title': 16, 'row': 24, 'margin': 0.75, 'spacing': 12}
        }
        self.sizes = density_map.get(args.density, density_map['normal'])
        
        # Override font size if specified
        if args.font_size:
            self.sizes['font'] = args.font_size
            self.sizes['label'] = args.font_size
        
        self.primary_rgb = RGBColor(
            int(self.primary_color[1:3], 16),
            int(self.primary_color[3:5], 16),
            int(self.primary_color[5:7], 16)
        )


# =============================================================================
# TEMPLATE GENERATION
# =============================================================================

class TemplateGenerator:
    """Main template generator"""
    
    def __init__(self, config, form_name, fields):
        self.config = config
        self.form_name = form_name
        self.fields = fields
        self.doc = Document()
        self.repeating_groups = detect_repeating_groups(fields)
        
        self._setup_document()
    
    def _setup_document(self):
        """Configure document settings"""
        section = self.doc.sections[0]
        
        # Orientation
        if self.config.orientation == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        
        # Margins
        margin = Inches(self.config.sizes['margin'])
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
        
        # Calculate usable width
        self.page_width = float(section.page_width - section.left_margin - section.right_margin) / 914400  # Convert to inches
    
    def _add_header_cell(self, cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
        """Add styled header cell"""
        set_cell_shading(cell, self.config.primary_color)
        set_cell_border(cell, self.config.primary_color.replace('#', ''))
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.font.size = Pt(self.config.sizes['label'])
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    def _add_section_header(self, text, color=None):
        """Add section header bar"""
        color = color or self.config.secondary_color
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(self.page_width)
        cell = table.rows[0].cells[0]
        set_cell_shading(cell, color)
        set_cell_border(cell, color.replace('#', ''))
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(self.config.sizes['label'])
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_row_height(table.rows[0], self.config.sizes['row'] - 4)
    
    def _add_label_value_row(self, table, row_idx, label, merge_field):
        """Add a label-value pair to a table row"""
        row = table.rows[row_idx]
        
        # Label cell
        cell = row.cells[0]
        set_cell_shading(cell, "F0F0F0")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(self.config.sizes['label'])
        run.font.name = 'Arial'
        run.font.bold = True
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Value cell
        cell = row.cells[1]
        set_cell_border(cell)
        p = cell.paragraphs[0]
        add_merge_field(p, merge_field, font_size=self.config.sizes['font'])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    def _add_spacer(self):
        """Add spacing paragraph"""
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(self.config.sizes['spacing'])
        spacer.paragraph_format.space_after = Pt(0)
    
    def generate(self):
        """Generate the full template"""
        self._add_title_banner()
        self._add_info_section()
        self._add_spacer()
        
        # Handle fields based on repeat-fields mode
        if self.repeating_groups and self.config.repeat_fields == 'grid':
            self._add_fields_as_grid()
        else:
            self._add_fields_stacked()
        
        self._add_spacer()
        self._add_signature_section()
        self._add_footer()
        
        return self.doc
    
    def _add_title_banner(self):
        """Add title banner with logo"""
        table = self.doc.add_table(rows=1, cols=3)
        table.autofit = False
        
        # Column widths
        logo_width = 1.2
        ref_width = 1.5
        title_width = self.page_width - logo_width - ref_width
        table.columns[0].width = Inches(title_width)
        table.columns[1].width = Inches(ref_width)
        table.columns[2].width = Inches(logo_width)
        set_row_height(table.rows[0], 36)
        
        # Title
        cell = table.rows[0].cells[0]
        set_cell_shading(cell, self.config.primary_color)
        set_cell_border(cell, self.config.primary_color.replace('#', ''))
        p = cell.paragraphs[0]
        # Clean up form name
        title = re.sub(r'\s*-\s*Thames Boilers$', '', self.form_name)
        title = title.upper()
        run = p.add_run(title)
        run.font.size = Pt(self.config.sizes['title'])
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Reference
        cell = table.rows[0].cells[1]
        set_cell_shading(cell, self.config.primary_color)
        set_cell_border(cell, self.config.primary_color.replace('#', ''))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("REF: ")
        run.font.size = Pt(self.config.sizes['font'] + 1)
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        add_merge_field(p, "job.generated_job_id", font_size=self.config.sizes['font'] + 1, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Logo
        cell = table.rows[0].cells[2]
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self.config.logo and os.path.exists(self.config.logo):
            run = p.add_run()
            run.add_picture(self.config.logo, width=Inches(1.0))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    def _add_info_section(self):
        """Add contractor/installation/client info"""
        # Headers
        header_table = self.doc.add_table(rows=1, cols=3)
        header_table.autofit = False
        col_width = self.page_width / 3
        for i in range(3):
            header_table.columns[i].width = Inches(col_width)
        
        headers = ["CONTRACTOR DETAILS", "INSTALLATION ADDRESS", "CLIENT DETAILS"]
        for i, header in enumerate(headers):
            self._add_header_cell(header_table.rows[0].cells[i], header)
        set_row_height(header_table.rows[0], self.config.sizes['row'] - 4)
        
        # Content
        content_table = self.doc.add_table(rows=1, cols=3)
        content_table.autofit = False
        for i in range(3):
            content_table.columns[i].width = Inches(col_width)
        
        fs = self.config.sizes['font']
        
        # Contractor
        cell = content_table.rows[0].cells[0]
        set_cell_border(cell)
        p = cell.paragraphs[0]
        add_merge_field(p, "vendor.name", font_size=fs, bold=True)
        p.add_run("\n")
        add_merge_field(p, "location.line1", font_size=fs-1)
        p.add_run(", ")
        add_merge_field(p, "location.state", font_size=fs-1)
        p.add_run(" ")
        add_merge_field(p, "location.post_code", font_size=fs-1)
        p.add_run("\nTel: ")
        add_merge_field(p, "location.phone_1", font_size=fs-1)
        
        # Installation
        cell = content_table.rows[0].cells[1]
        set_cell_border(cell)
        p = cell.paragraphs[0]
        add_merge_field(p, "job.contact_first", font_size=fs)
        p.add_run(" ")
        add_merge_field(p, "job.contact_last", font_size=fs)
        p.add_run("\n")
        add_merge_field(p, "job.job_address", font_size=fs-1)
        
        # Client
        cell = content_table.rows[0].cells[2]
        set_cell_border(cell)
        p = cell.paragraphs[0]
        add_merge_field(p, "job.company_name", font_size=fs)
        p.add_run("\n")
        add_merge_field(p, "job.billing_address", font_size=fs-1)
    
    def _add_fields_as_grid(self):
        """Add repeating fields as compact grid tables"""
        processed_fields = set()
        
        for group_name, group_data in self.repeating_groups.items():
            self._add_spacer()
            self._add_section_header(group_name.upper() + " DETAILS")
            
            # Get field bases (without numbers)
            field_bases = {}
            for field in group_data['fields']:
                match = re.match(r'^(.+?)[\s_]?\d+$', field['name'])
                if match:
                    base = match.group(1).strip()
                    if base not in field_bases:
                        field_bases[base] = field['type']
                    processed_fields.add(field['uuid'])
            
            # Separate single-line and multi-line fields
            single_line = {k: v for k, v in field_bases.items() if v not in ['Text (Multi-Line)', 'Signature']}
            multi_line = {k: v for k, v in field_bases.items() if v in ['Text (Multi-Line)']}
            signatures = {k: v for k, v in field_bases.items() if v == 'Signature'}
            
            # Create grid for single-line fields
            if single_line:
                num_cols = len(single_line) + 1  # +1 for row number
                num_rows = max(group_data['numbers']) + 1  # +1 for header
                
                grid = self.doc.add_table(rows=num_rows, cols=num_cols)
                grid.autofit = False
                
                # Calculate column widths
                num_width = 0.3
                remaining = self.page_width - num_width
                field_width = remaining / len(single_line)
                grid.columns[0].width = Inches(num_width)
                for i in range(1, num_cols):
                    grid.columns[i].width = Inches(field_width)
                
                # Header row
                header_row = grid.rows[0]
                set_row_height(header_row, self.config.sizes['row'] - 4)
                
                cell = header_row.cells[0]
                set_cell_shading(cell, "E0E0E0")
                set_cell_border(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("#")
                run.font.size = Pt(self.config.sizes['label'])
                run.font.name = 'Arial'
                run.font.bold = True
                
                for i, field_base in enumerate(single_line.keys()):
                    cell = header_row.cells[i + 1]
                    set_cell_shading(cell, "E0E0E0")
                    set_cell_border(cell)
                    p = cell.paragraphs[0]
                    # Shorten common labels
                    label = field_base.upper()
                    label = label.replace('APPLIANCE ', '').replace('CLASSIFICATION OF ', '').replace('NOTIFIABLE VIA ', '')
                    run = p.add_run(label)
                    run.font.size = Pt(self.config.sizes['label'])
                    run.font.name = 'Arial'
                    run.font.bold = True
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                # Data rows
                for row_num in sorted(group_data['numbers']):
                    row = grid.rows[row_num]
                    set_row_height(row, self.config.sizes['row'])
                    
                    # Row number
                    cell = row.cells[0]
                    set_cell_border(cell)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(str(row_num))
                    run.font.size = Pt(self.config.sizes['font'])
                    run.font.name = 'Arial'
                    run.font.bold = True
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    
                    # Field values
                    for i, field_base in enumerate(single_line.keys()):
                        cell = row.cells[i + 1]
                        set_cell_border(cell)
                        p = cell.paragraphs[0]
                        merge_name = name_to_merge_field(f"{field_base} {row_num}")
                        add_merge_field(p, merge_name, font_size=self.config.sizes['font'])
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # Create separate table for multi-line fields (notes)
            if multi_line:
                self._add_spacer()
                for field_base in multi_line.keys():
                    self._add_section_header(field_base.upper())
                    
                    notes_table = self.doc.add_table(rows=max(group_data['numbers']), cols=2)
                    notes_table.autofit = False
                    notes_table.columns[0].width = Inches(0.3)
                    notes_table.columns[1].width = Inches(self.page_width - 0.3)
                    
                    for row_num in sorted(group_data['numbers']):
                        row = notes_table.rows[row_num - 1]
                        set_row_height(row, self.config.sizes['row'] + 6)
                        
                        # Number
                        cell = row.cells[0]
                        set_cell_border(cell)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(str(row_num))
                        run.font.size = Pt(self.config.sizes['font'])
                        run.font.name = 'Arial'
                        run.font.bold = True
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                        
                        # Notes field
                        cell = row.cells[1]
                        set_cell_border(cell)
                        p = cell.paragraphs[0]
                        merge_name = name_to_merge_field(f"{field_base} {row_num}")
                        add_merge_field(p, merge_name, font_size=self.config.sizes['font'])
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        
        # Add non-repeating fields
        non_repeating = [f for f in self.fields if f['uuid'] not in processed_fields]
        if non_repeating:
            # Filter out signatures (handled separately)
            regular_fields = [f for f in non_repeating if f['type'] != 'Signature']
            if regular_fields:
                self._add_spacer()
                self._add_section_header("ADDITIONAL INFORMATION")
                self._add_fields_table(regular_fields)
    
    def _add_fields_stacked(self):
        """Add fields in traditional stacked layout"""
        self._add_section_header("FORM DETAILS")
        
        # Filter out signatures
        regular_fields = [f for f in self.fields if f['type'] != 'Signature']
        self._add_fields_table(regular_fields)
    
    def _add_fields_table(self, fields):
        """Add a simple two-column label-value table"""
        if not fields:
            return
        
        table = self.doc.add_table(rows=len(fields), cols=2)
        table.autofit = False
        table.columns[0].width = Inches(self.page_width * 0.35)
        table.columns[1].width = Inches(self.page_width * 0.65)
        
        for i, field in enumerate(fields):
            row = table.rows[i]
            row_height = self.config.sizes['row']
            if field['type'] == 'Text (Multi-Line)':
                row_height = self.config.sizes['row'] * 2
            set_row_height(row, row_height)
            
            # Label
            cell = row.cells[0]
            set_cell_shading(cell, "F5F5F5")
            set_cell_border(cell)
            p = cell.paragraphs[0]
            run = p.add_run(field['name'].upper())
            run.font.size = Pt(self.config.sizes['label'])
            run.font.name = 'Arial'
            run.font.bold = True
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # Value
            cell = row.cells[1]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            add_merge_field(p, field['merge_field'], font_size=self.config.sizes['font'])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    def _add_signature_section(self):
        """Add signature section"""
        self._add_section_header("DECLARATION", self.config.primary_color)
        
        # Find signature fields
        sig_fields = [f for f in self.fields if f['type'] == 'Signature']
        
        # Standard signature layout
        sig_table = self.doc.add_table(rows=2, cols=8)
        sig_table.autofit = False
        
        col_width = self.page_width / 8
        for i in range(8):
            sig_table.columns[i].width = Inches(col_width)
        
        # Engineer row
        row = sig_table.rows[0]
        set_row_height(row, 28)
        
        labels = [("ENGINEER", "staff.full_name"), ("SIGNED", "form_engineer_signature"), 
                  ("ID NO.", "form_gas_safe_card_number"), ("DATE", "calculation.todays_date")]
        
        for i, (label, field) in enumerate(labels):
            # Label
            cell = row.cells[i * 2]
            set_cell_shading(cell, "F0F0F0")
            set_cell_border(cell)
            p = cell.paragraphs[0]
            run = p.add_run(label)
            run.font.size = Pt(self.config.sizes['label'])
            run.font.name = 'Arial'
            run.font.bold = True
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # Value
            cell = row.cells[i * 2 + 1]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            add_merge_field(p, field, font_size=self.config.sizes['font'])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Customer row
        row = sig_table.rows[1]
        set_row_height(row, 28)
        
        labels = [("CUSTOMER", "form_customer_name"), ("SIGNED", "form_customer_signature"),
                  ("PRESENT", "form_customer_present"), ("", "")]
        
        for i, (label, field) in enumerate(labels):
            # Label
            cell = row.cells[i * 2]
            if label:
                set_cell_shading(cell, "F0F0F0")
            set_cell_border(cell)
            p = cell.paragraphs[0]
            if label:
                run = p.add_run(label)
                run.font.size = Pt(self.config.sizes['label'])
                run.font.name = 'Arial'
                run.font.bold = True
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # Value
            cell = row.cells[i * 2 + 1]
            set_cell_border(cell)
            if field:
                p = cell.paragraphs[0]
                add_merge_field(p, field, font_size=self.config.sizes['font'])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    def _add_footer(self):
        """Add footer"""
        self._add_spacer()
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"{self.config.company}")
        run.font.size = Pt(self.config.sizes['font'] - 1)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(128, 128, 128)


# =============================================================================
# MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description='Generate ServiceM8 form templates')
    parser.add_argument('--sm8f', required=True, help='Path to SM8F form file')
    parser.add_argument('--company', required=True, help='Company name')
    parser.add_argument('--logo', help='Path to logo image')
    parser.add_argument('--primary', default='#004aad', help='Primary brand color (hex)')
    parser.add_argument('--style', choices=['classic', 'bold_trade', 'modern_corporate', 'minimal'],
                        default='bold_trade', help='Visual style')
    parser.add_argument('--layout', choices=['compact', 'expanded'], default='compact',
                        help='Layout density')
    parser.add_argument('--density', choices=['tight', 'normal', 'spacious'], default='tight',
                        help='Content density')
    parser.add_argument('--repeat-fields', choices=['grid', 'stacked'], default='grid',
                        help='How to display repeating field groups')
    parser.add_argument('--orientation', choices=['portrait', 'landscape'], default='portrait',
                        help='Page orientation')
    parser.add_argument('--font-size', type=int, help='Override base font size (pt)')
    parser.add_argument('--output', required=True, help='Output DOCX path')
    
    args = parser.parse_args()
    
    # Parse form
    form_name, fields = extract_form(args.sm8f)
    print(f"📋 Form: {form_name}")
    print(f"   Fields: {len(fields)}")
    
    # Create config
    config = Config(args)
    
    # Detect repeating groups
    groups = detect_repeating_groups(fields)
    if groups:
        print(f"   Repeating groups: {', '.join(groups.keys())}")
    
    # Generate
    generator = TemplateGenerator(config, form_name, fields)
    doc = generator.generate()
    doc.save(args.output)
    
    print(f"\n✅ Template generated: {args.output}")
    print(f"   Style: {args.style}")
    print(f"   Layout: {args.layout} / {args.density}")
    print(f"   Orientation: {args.orientation}")
    print(f"   Repeat fields: {args.repeat_fields}")

if __name__ == "__main__":
    main()
